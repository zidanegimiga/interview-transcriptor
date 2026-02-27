import asyncio
from datetime import datetime, timezone

from bson import ObjectId
from fastapi import APIRouter, HTTPException, UploadFile, File, Form, status

from app.core.deps import CurrentUser, DBDep
from app.core.config import settings
from app.models.common import ok, paginate
from app.models.interview import InterviewStatus, UpdateInterviewRequest
from app.services.storage import get_storage_backend
from app.services.transcription import get_transcription_service
from app.services.analysis import run_analysis

router = APIRouter(prefix="/interviews", tags=["Interviews"])


def _serialize(doc: dict) -> dict:
    """Convert ObjectId fields to strings for JSON serialisation."""
    doc["_id"] = str(doc["_id"])
    doc["user_id"] = str(doc["user_id"])
    if doc.get("template_id"):
        doc["template_id"] = str(doc["template_id"])
    return doc


# Upload 

@router.post("/upload", status_code=status.HTTP_201_CREATED)
async def upload_interview(
    user: CurrentUser,
    db: DBDep,
    file: UploadFile = File(...),
    title: str | None = Form(None),
):
    # Validate file type
    if file.content_type not in settings.ALLOWED_MIME_TYPES:
        raise HTTPException(
            status_code=status.HTTP_415_UNSUPPORTED_MEDIA_TYPE,
            detail=f"File type '{file.content_type}' is not supported. "
                   f"Allowed types: {', '.join(settings.ALLOWED_MIME_TYPES)}",
        )

    # Read and validate file size
    file_bytes = await file.read()
    if len(file_bytes) > settings.MAX_FILE_SIZE_BYTES:
        raise HTTPException(
            status_code=status.HTTP_413_REQUEST_ENTITY_TOO_LARGE,
            detail=f"File exceeds maximum size of {settings.MAX_FILE_SIZE_MB}MB.",
        )

    # Upload to R2
    storage = get_storage_backend()
    try:
        storage_key = await storage.upload(
            file_bytes=file_bytes,
            filename=file.filename,
            content_type=file.content_type,
        )
    except Exception:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail="File upload failed. Please try again.",
        )

    # Create MongoDB document
    now = datetime.now(timezone.utc)
    document = {
        "user_id": user["id"],
        "title": title or file.filename,
        "original_name": file.filename,
        "filename": storage_key.split("/")[-1],
        "file_size": len(file_bytes),
        "file_type": file.content_type,
        "storage_key": storage_key,
        "duration_seconds": None,
        "status": InterviewStatus.uploaded.value,
        "error_message": None,
        "deepgram_job_id": None,
        "template_id": None,
        "transcript": None,
        "ai_analysis": None,
        "tags": [],
        "notifications": [],
        "created_at": now,
        "updated_at": now,
    }

    result = await db["interviews"].insert_one(document)
    document["_id"] = str(result.inserted_id)
    document["user_id"] = str(document["user_id"])

    return ok(document)


# List 

@router.get("")
async def list_interviews(
    user: CurrentUser,
    db: DBDep,
    page: int = 1,
    limit: int = 20,
    interview_status: str | None = None,
    search: str | None = None,
):
    query: dict = {"user_id": user["id"]}
    if interview_status:
        query["status"] = interview_status
    if search:
        query["$text"] = {"$search": search}

    skip = (page - 1) * limit
    total = await db["interviews"].count_documents(query)
    cursor = db["interviews"].find(query).sort("created_at", -1).skip(skip).limit(limit)
    docs = await cursor.to_list(length=limit)

    return ok([_serialize(d) for d in docs], paginate(page, limit, total))


# Metrics 

@router.get("/metrics")
async def get_metrics(user: CurrentUser, db: DBDep):
    status_pipeline = [
        {"$match": {"user_id": user["id"]}},
        {"$group": {"_id": "$status", "count": {"$sum": 1}}},
    ]
    status_counts = await db["interviews"].aggregate(status_pipeline).to_list(length=20)

    sentiment_pipeline = [
        {"$match": {"user_id": user["id"], "ai_analysis": {"$ne": None}}},
        {"$group": {"_id": "$ai_analysis.sentiment.overall", "count": {"$sum": 1}}},
    ]
    sentiment_counts = await db["interviews"].aggregate(sentiment_pipeline).to_list(length=10)

    keyword_pipeline = [
        {"$match": {"user_id": user["id"], "ai_analysis.keywords": {"$exists": True}}},
        {"$unwind": "$ai_analysis.keywords"},
        {"$group": {"_id": "$ai_analysis.keywords.term", "count": {"$sum": "$ai_analysis.keywords.frequency"}}},
        {"$sort": {"count": -1}},
        {"$limit": 10},
    ]
    top_keywords = await db["interviews"].aggregate(keyword_pipeline).to_list(length=10)

    return ok({
        "by_status": {s["_id"]: s["count"] for s in status_counts},
        "by_sentiment": {s["_id"]: s["count"] for s in sentiment_counts},
        "top_keywords": [{"term": k["_id"], "count": k["count"]} for k in top_keywords],
    })


# Get one 

@router.get("/{interview_id}")
async def get_interview(interview_id: str, user: CurrentUser, db: DBDep):
    try:
        oid = ObjectId(interview_id)
    except Exception:
        raise HTTPException(status_code=400, detail="Invalid interview ID.")

    doc = await db["interviews"].find_one({"_id": oid, "user_id": user["id"]})
    if not doc:
        raise HTTPException(status_code=404, detail="Interview not found.")

    return ok(_serialize(doc))


# Update
@router.patch("/{interview_id}")
async def update_interview(
    interview_id: str,
    payload: UpdateInterviewRequest,
    user: CurrentUser,
    db: DBDep,
):
    try:
        oid = ObjectId(interview_id)
    except Exception:
        raise HTTPException(status_code=400, detail="Invalid interview ID.")

    updates = payload.model_dump(exclude_none=True)
    if not updates:
        raise HTTPException(status_code=400, detail="No fields provided to update.")
    updates["updated_at"] = datetime.now(timezone.utc)

    result = await db["interviews"].update_one(
        {"_id": oid, "user_id": user["id"]},
        {"$set": updates},
    )
    if result.matched_count == 0:
        raise HTTPException(status_code=404, detail="Interview not found.")

    doc = await db["interviews"].find_one({"_id": oid})
    return ok(_serialize(doc))


# Delete

@router.delete("/{interview_id}", status_code=status.HTTP_204_NO_CONTENT)
async def delete_interview(interview_id: str, user: CurrentUser, db: DBDep):
    try:
        oid = ObjectId(interview_id)
    except Exception:
        raise HTTPException(status_code=400, detail="Invalid interview ID.")

    result = await db["interviews"].delete_one({"_id": oid, "user_id": user["id"]})
    if result.deleted_count == 0:
        raise HTTPException(status_code=404, detail="Interview not found.")


# ── Status 

@router.get("/{interview_id}/status")
async def get_status(interview_id: str, user: CurrentUser, db: DBDep):
    try:
        oid = ObjectId(interview_id)
    except Exception:
        raise HTTPException(status_code=400, detail="Invalid interview ID.")

    doc = await db["interviews"].find_one(
        {"_id": oid, "user_id": user["id"]},
        {"status": 1, "error_message": 1, "updated_at": 1},
    )
    if not doc:
        raise HTTPException(status_code=404, detail="Interview not found.")

    return ok({
        "id": interview_id,
        "status": doc["status"],
        "error_message": doc.get("error_message"),
        "updated_at": doc["updated_at"].isoformat(),
    })


# transcribe 

@router.post("/{interview_id}/transcribe")
async def transcribe_interview(interview_id: str, user: CurrentUser, db: DBDep):
    try:
        oid = ObjectId(interview_id)
    except Exception:
        raise HTTPException(status_code=400, detail="Invalid interview ID.")

    doc = await db["interviews"].find_one({"_id": oid, "user_id": user["id"]})
    if not doc:
        raise HTTPException(status_code=404, detail="Interview not found.")

    if doc.get("status") in ("transcribing", "analysing", "completed"):
        return ok({"id": interview_id, "message": "Already transcribed or in progress."})

    try:
        job_id = await get_transcription_service().submit(
            storage_key=doc["storage_key"],
            interview_id=interview_id,
        )
    except Exception as exc:
        raise HTTPException(
            status_code=500,
            detail=f"Transcription submission failed: {str(exc)}"
        )

    await db["interviews"].update_one(
        {"_id": oid},
        {"$set": {
            "status":          "queued",
            "deepgram_job_id": job_id,
            "updated_at":      datetime.now(timezone.utc),
        }}
    )

    return ok({"id": interview_id, "status": "queued", "job_id": job_id})


# Analyse
@router.post("/{interview_id}/analyse")
async def analyse_interview(interview_id: str, user: CurrentUser, db: DBDep):
    try:
        oid = ObjectId(interview_id)
    except Exception:
        raise HTTPException(status_code=400, detail="Invalid interview ID.")

    doc = await db["interviews"].find_one({"_id": oid, "user_id": user["id"]})
    if not doc:
        raise HTTPException(status_code=404, detail="Interview not found.")

    if doc.get("ai_analysis"):
        return ok({"id": interview_id, "message": "Already analysed."})

    if not doc.get("transcript"):
        raise HTTPException(
            status_code=400,
            detail="Transcript not available. Run transcription first."
        )

    await db["interviews"].update_one(
        {"_id": oid},
        {"$set": {"status": "analysing", "updated_at": datetime.now(timezone.utc)}}
    )

    asyncio.create_task(run_analysis(interview_id, user["id"]))

    return ok({"id": interview_id, "message": "Analysis started."})


# Export
@router.get("/{interview_id}/export")
async def export_transcript(
    interview_id: str,
    user: CurrentUser,
    db: DBDep,
    format: str = "txt",
):
    try:
        oid = ObjectId(interview_id)
    except Exception:
        raise HTTPException(status_code=400, detail="Invalid interview ID.")

    doc = await db["interviews"].find_one({"_id": oid, "user_id": user["id"]})
    if not doc:
        raise HTTPException(status_code=404, detail="Interview not found.")

    transcript = doc.get("transcript")
    if not transcript or not transcript.get("text"):
        raise HTTPException(status_code=400, detail="No transcript available to export.")

    title = doc.get("title", "Interview")
    transcript_text = transcript["text"]
    utterances = transcript.get("utterances", [])

    if format == "txt":
        return _export_txt(title, utterances, transcript_text)
    elif format == "pdf":
        return _export_pdf(title, utterances, transcript_text)
    elif format == "docx":
        return _export_docx(title, utterances, transcript_text)
    else:
        raise HTTPException(status_code=400, detail="Invalid format. Use txt, pdf, or docx.")


# Batch upload
@router.post("/batch-upload", status_code=status.HTTP_201_CREATED)
async def batch_upload(
    user: CurrentUser,
    db: DBDep,
    files: list[UploadFile] = File(...),
):
    if len(files) > settings.MAX_BATCH_FILES:
        raise HTTPException(
            status_code=400,
            detail=f"Maximum {settings.MAX_BATCH_FILES} files per batch.",
        )

    storage = get_storage_backend()
    created = []
    failed = []
    now = datetime.now(timezone.utc)

    for file in files:
        # Validate MIME type
        if file.content_type not in settings.ALLOWED_MIME_TYPES:
            failed.append({
                "filename": file.filename,
                "error": f"Unsupported file type: {file.content_type}",
            })
            continue

        # Read and validate size
        file_bytes = await file.read()
        if len(file_bytes) > settings.MAX_FILE_SIZE_BYTES:
            failed.append({
                "filename": file.filename,
                "error": f"File exceeds {settings.MAX_FILE_SIZE_MB}MB limit.",
            })
            continue

        # Upload to R2
        try:
            storage_key = await storage.upload(
                file_bytes=file_bytes,
                filename=file.filename,
                content_type=file.content_type,
            )
        except Exception as exc:
            failed.append({
                "filename": file.filename,
                "error": "Storage upload failed.",
            })
            continue

        # Create MongoDB document
        document = {
            "user_id":        user["id"],
            "title":          file.filename,
            "original_name":  file.filename,
            "filename":       storage_key.split("/")[-1],
            "file_size":      len(file_bytes),
            "file_type":      file.content_type,
            "storage_key":    storage_key,
            "duration_seconds": None,
            "status":         InterviewStatus.uploaded.value,
            "error_message":  None,
            "deepgram_job_id": None,
            "template_id":    None,
            "transcript":     None,
            "ai_analysis":    None,
            "tags":           [],
            "notifications":  [],
            "created_at":     now,
            "updated_at":     now,
        }

        result = await db["interviews"].insert_one(document)
        created.append({
            "id":       str(result.inserted_id),
            "filename": file.filename,
            "status":   InterviewStatus.uploaded.value,
        })

    return ok({
        "created": created,
        "failed":  failed,
        "total_created": len(created),
        "total_failed":  len(failed),
    })




# ── Export helpers

def _format_utterances(utterances: list, transcript_text: str) -> str:
    """Format utterances as Speaker A: ... Speaker B: ... or fall back to raw text."""
    if not utterances:
        return transcript_text
    lines = []
    for u in utterances:
        speaker = u.get("speaker", "Unknown")
        text = u.get("text", "")
        lines.append(f"Speaker {speaker}: {text}")
    return "\n\n".join(lines)


def _export_txt(title: str, utterances: list, transcript_text: str):
    from fastapi.responses import PlainTextResponse
    content = f"{title}\n{'=' * len(title)}\n\n"
    content += _format_utterances(utterances, transcript_text)
    return PlainTextResponse(
        content=content,
        headers={"Content-Disposition": f'attachment; filename="{title}.txt"'},
    )


def _export_pdf(title: str, utterances: list, transcript_text: str):
    import io
    from fastapi.responses import StreamingResponse
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer

    buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        leftMargin=2.5 * cm,
        rightMargin=2.5 * cm,
        topMargin=2.5 * cm,
        bottomMargin=2.5 * cm,
    )

    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "Title",
        parent=styles["Heading1"],
        fontSize=16,
        spaceAfter=20,
    )
    speaker_style = ParagraphStyle(
        "Speaker",
        parent=styles["Normal"],
        fontSize=10,
        fontName="Helvetica-Bold",
        spaceAfter=4,
    )
    text_style = ParagraphStyle(
        "Text",
        parent=styles["Normal"],
        fontSize=10,
        spaceAfter=12,
        leading=14,
    )

    story = [Paragraph(title, title_style), Spacer(1, 0.5 * cm)]

    formatted = _format_utterances(utterances, transcript_text)
    for block in formatted.split("\n\n"):
        if block.startswith("Speaker "):
            parts = block.split(": ", 1)
            if len(parts) == 2:
                story.append(Paragraph(parts[0] + ":", speaker_style))
                story.append(Paragraph(parts[1], text_style))
            else:
                story.append(Paragraph(block, text_style))
        else:
            story.append(Paragraph(block, text_style))

    doc.build(story)
    buffer.seek(0)

    return StreamingResponse(
        buffer,
        media_type="application/pdf",
        headers={"Content-Disposition": f'attachment; filename="{title}.pdf"'},
    )


def _export_docx(title: str, utterances: list, transcript_text: str):
    import io
    from fastapi.responses import StreamingResponse
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    document = Document()

    # Title
    heading = document.add_heading(title, level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT

    document.add_paragraph()

    formatted = _format_utterances(utterances, transcript_text)
    for block in formatted.split("\n\n"):
        if block.startswith("Speaker "):
            parts = block.split(": ", 1)
            if len(parts) == 2:
                # Speaker label in bold
                p = document.add_paragraph()
                run = p.add_run(parts[0] + ":")
                run.bold = True
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)
                # Text on next paragraph
                text_p = document.add_paragraph(parts[1])
                text_p.paragraph_format.space_after = Pt(8)
            else:
                document.add_paragraph(block)
        else:
            document.add_paragraph(block)

    buffer = io.BytesIO()
    document.save(buffer)
    buffer.seek(0)

    return StreamingResponse(
        buffer,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{title}.docx"'},
    )